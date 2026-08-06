"""
Text extraction pipeline — Phase 2.

Handles: PDF (digital + scanned detection), DOCX, DOC (via LibreOffice), Images.

Architecture
------------
extract_document(path, mime_type) -> ExtractionResult
  └── _extract_pdf()        — pdfplumber with per-page column detection
      └── _extract_page_digital()   — coordinate-aware, reconstructs two-column
      └── _extract_page_scanned()   — OCR placeholder (flagged for Phase 3)
  └── _extract_docx()       — python-docx, tables → key-value pairs
  └── _extract_doc()        — LibreOffice headless → DOCX → python-docx
  └── _extract_image()      — OpenCV pre-process → OCR placeholder
  └── _clean_text()         — ftfy, page markers, repeated headers, bullets
"""

from __future__ import annotations

import logging
import os
import re
import shutil
import subprocess
import tempfile
import threading
import zipfile
from dataclasses import dataclass, field
from pathlib import Path
from typing import Optional

import ftfy

logger = logging.getLogger(__name__)

# ---------------------------------------------------------------------------
# Result type
# ---------------------------------------------------------------------------

@dataclass
class ExtractionResult:
    text: str                          # clean extracted text
    page_count: int = 0
    scanned_page_count: int = 0        # pages that had no digital text layer
    ocr_required: bool = False         # true if any page was scanned
    extraction_method: str = "unknown" # pdf_digital | pdf_mixed | docx | doc | image
    warnings: list[str] = field(default_factory=list)
    confidence: float = 1.0            # 0–1; lower if many scanned pages or errors


@dataclass
class OcrResult:
    text: str
    engine: str = "none"
    confidence: float = 0.0
    warning: Optional[str] = None


_PADDLE_LOCK = threading.Lock()
_PADDLE_ENGINE = None
_PADDLE_INIT_ERROR: Optional[str] = None


# ---------------------------------------------------------------------------
# Text cleaning helpers
# ---------------------------------------------------------------------------

# Matches page markers like:  -- 1 of 5 --   |  Page 2 of 3  |  - 1 -
_PAGE_MARKER_RE = re.compile(
    r"(?m)^[-\s]*(?:Page\s+)?\d+\s+(?:of\s+\d+\s*)?[-\s]*$",
    re.IGNORECASE,
)
# Also catches exact patterns from real CVs:  "-- 2 of 5 --"  "-- 4 of 4 --"
_DASH_PAGE_RE = re.compile(r"(?m)^--\s*\d+\s+of\s+\d+\s*--\s*$")
# Separator lines of dashes/equals (artifacts)
_SEPARATOR_RE = re.compile(r"(?m)^[-=_]{10,}\s*$")
# Unicode bullet variants → standard hyphen-space
_BULLET_RE = re.compile(r"[•●◆▪·∙▸►➤➢◉○→]")
# Multiple blank lines → max two
_MULTI_BLANK_RE = re.compile(r"\n{3,}")


def _clean_text(raw: str, page_texts: Optional[list[str]] = None) -> str:
    """
    Full cleaning pipeline:
    1. ftfy  — fix broken Unicode/encoding
    2. Remove repeated page headers/footers (lines appearing on 3+ pages)
    3. Remove page markers
    4. Remove separator lines
    5. Normalize bullets
    6. Normalize whitespace and blank lines
    """
    if not raw or not raw.strip():
        return ""

    # 1. Fix encoding/Unicode issues (LinkedIn icons become ✓ / stripped etc.)
    text = ftfy.fix_text(raw)

    # 2. Remove repeated headers/footers across pages
    if page_texts and len(page_texts) >= 3:
        text = _remove_repeated_headers(text, page_texts)

    # 3. Page markers
    text = _DASH_PAGE_RE.sub("", text)
    text = _PAGE_MARKER_RE.sub("", text)

    # 4. Separator lines
    text = _SEPARATOR_RE.sub("", text)

    # 5. Normalize bullets
    text = _BULLET_RE.sub("- ", text)

    # 6. Whitespace normalization
    # Join lines split mid-sentence (no punctuation at end, next starts lowercase)
    text = _rejoin_broken_lines(text)
    # Collapse multiple blank lines
    text = _MULTI_BLANK_RE.sub("\n\n", text)
    # Strip trailing whitespace on each line
    text = "\n".join(line.rstrip() for line in text.splitlines())

    return text.strip()


def _remove_repeated_headers(full_text: str, page_texts: list[str]) -> str:
    """
    Find lines that appear verbatim on 3+ pages — these are headers/footers.
    Remove them from the final text.
    """
    if len(page_texts) < 3:
        return full_text

    # Count line frequency across pages
    line_freq: dict[str, int] = {}
    for page in page_texts:
        seen_on_this_page: set[str] = set()
        for line in page.splitlines():
            stripped = line.strip()
            if len(stripped) > 3 and stripped not in seen_on_this_page:
                line_freq[stripped] = line_freq.get(stripped, 0) + 1
                seen_on_this_page.add(stripped)

    # Lines appearing on ≥ 3 pages are repeated headers/footers
    repeated = {line for line, count in line_freq.items() if count >= 3}
    if not repeated:
        return full_text

    logger.debug("[extractor] removing %d repeated header/footer lines", len(repeated))
    cleaned_lines = []
    for line in full_text.splitlines():
        if line.strip() not in repeated:
            cleaned_lines.append(line)
    return "\n".join(cleaned_lines)


def _rejoin_broken_lines(text: str) -> str:
    """
    Join lines that are continuations of the previous line:
    - Previous line does NOT end with sentence-ending punctuation
    - Current line starts with a lowercase letter
    This fixes mid-sentence line breaks common in two-column PDF extractions.
    """
    lines = text.splitlines()
    result: list[str] = []
    i = 0
    while i < len(lines):
        line = lines[i]
        # While next line looks like a continuation, merge
        while (
            i + 1 < len(lines)
            and line.strip()
            and lines[i + 1].strip()
            and not re.search(r"[.!?:;,]\s*$", line.rstrip())
            and re.match(r"^[a-z]", lines[i + 1].lstrip())
        ):
            line = line.rstrip() + " " + lines[i + 1].lstrip()
            i += 1
        result.append(line)
        i += 1
    return "\n".join(result)


# ---------------------------------------------------------------------------
# OCR engines
# ---------------------------------------------------------------------------

def _ocr_setting(name: str, default):
    """Read OCR config lazily so this service remains usable in isolation."""
    try:
        import config as app_config

        return getattr(app_config, name, default)
    except Exception:
        return default


def _paddle_language() -> str:
    languages = str(_ocr_setting("DOCUMENT_OCR_LANGUAGES", "en")).lower()
    primary = re.split(r"[,+\s]+", languages)[0]
    return {
        "eng": "en",
        "english": "en",
        "hin": "hi",
        "hindi": "hi",
    }.get(primary, primary or "en")


def _get_paddle_engine():
    global _PADDLE_ENGINE, _PADDLE_INIT_ERROR

    if _PADDLE_ENGINE is not None:
        return _PADDLE_ENGINE
    if _PADDLE_INIT_ERROR is not None:
        return None

    with _PADDLE_LOCK:
        if _PADDLE_ENGINE is not None:
            return _PADDLE_ENGINE
        if _PADDLE_INIT_ERROR is not None:
            return None
        try:
            from paddleocr import PaddleOCR

            _PADDLE_ENGINE = PaddleOCR(
                lang=_paddle_language(),
                use_doc_orientation_classify=True,
                use_doc_unwarping=False,
                use_textline_orientation=True,
            )
        except Exception as exc:
            _PADDLE_INIT_ERROR = str(exc)
            logger.info("[extractor] PaddleOCR unavailable: %s", exc)
        return _PADDLE_ENGINE


def _ocr_with_paddle(image_path: Path) -> OcrResult:
    engine = _get_paddle_engine()
    if engine is None:
        return OcrResult(
            text="",
            engine="paddle",
            warning=f"PaddleOCR unavailable: {_PADDLE_INIT_ERROR or 'not installed'}",
        )

    try:
        with _PADDLE_LOCK:
            results = engine.predict(input=str(image_path))

        texts: list[str] = []
        scores: list[float] = []
        for result in results:
            payload = getattr(result, "json", result)
            if callable(payload):
                payload = payload()
            if not isinstance(payload, dict):
                continue
            data = payload.get("res", payload)
            texts.extend(str(text).strip() for text in data.get("rec_texts", []) if str(text).strip())
            scores.extend(float(score) for score in data.get("rec_scores", []) if score is not None)

        text = _clean_text("\n".join(texts))
        confidence = sum(scores) / len(scores) if scores else (0.7 if text else 0.0)
        return OcrResult(text=text, engine="paddle", confidence=confidence)
    except Exception as exc:
        logger.warning("[extractor] PaddleOCR failed for %s: %s", image_path.name, exc)
        return OcrResult(text="", engine="paddle", warning=f"PaddleOCR failed: {exc}")


def _ocr_with_tesseract(image_path: Path) -> OcrResult:
    try:
        import pytesseract
        from PIL import Image

        command = str(_ocr_setting("TESSERACT_CMD", "")).strip()
        candidates = [
            command,
            shutil.which("tesseract") or "",
            str(Path(os.environ.get("LOCALAPPDATA", "")) / "Programs" / "Tesseract-OCR" / "tesseract.exe"),
            str(Path(os.environ.get("ProgramFiles", "")) / "Tesseract-OCR" / "tesseract.exe"),
        ]
        executable = next((path for path in candidates if path and Path(path).is_file()), "")
        if executable:
            pytesseract.pytesseract.tesseract_cmd = executable

        language = str(_ocr_setting("DOCUMENT_OCR_LANGUAGES", "eng")).strip() or "eng"
        with Image.open(image_path) as image:
            data = pytesseract.image_to_data(
                image,
                lang=language,
                output_type=pytesseract.Output.DICT,
            )

        lines: dict[tuple[int, int, int], list[str]] = {}
        confidences: list[float] = []
        count = len(data.get("text", []))
        for i in range(count):
            word = str(data["text"][i]).strip()
            try:
                confidence = float(data["conf"][i])
            except (TypeError, ValueError):
                confidence = -1.0
            if not word or confidence < 0:
                continue
            key = (
                int(data["block_num"][i]),
                int(data["par_num"][i]),
                int(data["line_num"][i]),
            )
            lines.setdefault(key, []).append(word)
            confidences.append(confidence / 100.0)

        text = _clean_text("\n".join(" ".join(words) for words in lines.values()))
        confidence = sum(confidences) / len(confidences) if confidences else 0.0
        return OcrResult(text=text, engine="tesseract", confidence=confidence)
    except ImportError:
        return OcrResult(text="", engine="tesseract", warning="pytesseract is not installed")
    except Exception as exc:
        logger.warning("[extractor] Tesseract failed for %s: %s", image_path.name, exc)
        return OcrResult(text="", engine="tesseract", warning=f"Tesseract failed: {exc}")


def _run_ocr(original_path: Path, preprocessed_path: Optional[Path] = None) -> OcrResult:
    """
    Run the configured OCR engine. `auto` prefers PaddleOCR and falls back to
    Tesseract. Paddle receives the original image so its own orientation and
    layout models can work; Tesseract receives the OpenCV-cleaned image.
    """
    requested = str(_ocr_setting("DOCUMENT_OCR_ENGINE", "auto")).strip().lower()
    warnings: list[str] = []

    if requested in {"auto", "paddle"}:
        result = _ocr_with_paddle(original_path)
        if result.text:
            return result
        if result.warning:
            warnings.append(result.warning)
        if requested == "paddle":
            result.warning = "; ".join(warnings)
            return result

    if requested in {"auto", "tesseract"}:
        result = _ocr_with_tesseract(preprocessed_path or original_path)
        if result.text:
            return result
        if result.warning:
            warnings.append(result.warning)
        result.warning = "; ".join(warnings) or "OCR returned no text"
        return result

    return OcrResult(
        text="",
        warning=f"Unknown DOCUMENT_OCR_ENGINE={requested!r}; expected auto, paddle, or tesseract",
    )


# ---------------------------------------------------------------------------
# PDF extraction — the most complex path
# ---------------------------------------------------------------------------

# Minimum character count to consider a page "digital" (has a text layer)
_MIN_DIGITAL_CHARS = 30
# Column gap threshold: if the x-coordinate spread of words spans two clusters
# separated by more than this fraction of page width, treat as two-column.
_COLUMN_GAP_RATIO = 0.12


def _detect_columns(words: list[dict], page_width: float) -> bool:
    """
    Return True if the page words form two distinct x-coordinate clusters.
    Uses a simple gap-detection approach on the left edge (x0) of words.
    """
    if not words or page_width == 0:
        return False
    x0_values = sorted(w["x0"] for w in words if w.get("text", "").strip())
    if len(x0_values) < 10:
        return False

    gap_threshold = page_width * _COLUMN_GAP_RATIO
    # Look for a gap > threshold in the middle 40–70% of x range
    x_min, x_max = x0_values[0], x0_values[-1]
    mid_start = x_min + (x_max - x_min) * 0.30
    mid_end   = x_min + (x_max - x_min) * 0.70

    mid_values = [x for x in x0_values if mid_start <= x <= mid_end]
    if not mid_values:
        return False

    # Check if there's a gap > threshold within mid_values
    mid_sorted = sorted(mid_values)
    for j in range(len(mid_sorted) - 1):
        if mid_sorted[j + 1] - mid_sorted[j] > gap_threshold:
            return True
    return False


def _extract_two_column_text(words: list[dict], page_width: float) -> str:
    """
    For two-column pages: split words into left/right columns, sort each
    column top-to-bottom, then concatenate left + right.
    """
    mid_x = page_width / 2
    left  = [w for w in words if w["x0"] < mid_x and w.get("text", "").strip()]
    right = [w for w in words if w["x0"] >= mid_x and w.get("text", "").strip()]

    def words_to_text(word_list: list[dict]) -> str:
        # Group into lines by y0 proximity (within 3 points = same line)
        if not word_list:
            return ""
        sorted_words = sorted(word_list, key=lambda w: (round(w["top"] / 3), w["x0"]))
        lines: list[str] = []
        current_line: list[str] = []
        current_y = sorted_words[0]["top"]
        for w in sorted_words:
            if abs(w["top"] - current_y) > 5:
                if current_line:
                    lines.append(" ".join(current_line))
                current_line = [w["text"]]
                current_y = w["top"]
            else:
                current_line.append(w["text"])
        if current_line:
            lines.append(" ".join(current_line))
        return "\n".join(lines)

    return words_to_text(left) + "\n\n" + words_to_text(right)


def _extract_page_digital(page) -> str:  # page is pdfplumber.Page
    """Extract text from a digital PDF page, handling two-column layouts."""
    try:
        words = page.extract_words(
            x_tolerance=3,
            y_tolerance=3,
            keep_blank_chars=False,
        )
        if not words:
            return ""

        if _detect_columns(words, page.width):
            return _extract_two_column_text(words, page.width)

        # Single-column: use standard extract_text (preserves layout)
        text = page.extract_text(x_tolerance=3, y_tolerance=3) or ""
        return text
    except Exception as exc:
        logger.warning("[extractor] page digital extraction error: %s", exc)
        return ""


def _ocr_pdf_page(pdf_document, page_index: int) -> OcrResult:
    """Render one PDF page to a temporary image and run OCR."""
    import fitz

    rendered_path: Optional[Path] = None
    preprocessed_path: Optional[Path] = None
    try:
        dpi = max(150, min(400, int(_ocr_setting("DOCUMENT_OCR_PDF_DPI", 250))))
        page = pdf_document.load_page(page_index)
        scale = dpi / 72.0
        pixmap = page.get_pixmap(matrix=fitz.Matrix(scale, scale), alpha=False)
        with tempfile.NamedTemporaryFile(suffix=".png", delete=False) as tmp:
            rendered_path = Path(tmp.name)
        pixmap.save(str(rendered_path))
        preprocessed_path = _preprocess_image_for_ocr(rendered_path)
        return _run_ocr(rendered_path, preprocessed_path)
    except Exception as exc:
        logger.warning("[extractor] PDF page %d OCR failed: %s", page_index + 1, exc)
        return OcrResult(text="", warning=f"Page {page_index + 1} OCR failed: {exc}")
    finally:
        for temp_path in (preprocessed_path, rendered_path):
            if temp_path:
                try:
                    temp_path.unlink(missing_ok=True)
                except OSError:
                    pass


def _extract_pdf(path: Path) -> ExtractionResult:
    """
    Extract text from a PDF.
    Per-page: detect if digital or scanned; handle two-column layouts.
    """
    import fitz
    import pdfplumber

    page_texts: list[str] = []
    scanned_pages: list[int] = []
    ocr_scores: list[float] = []
    ocr_engines: set[str] = set()
    warnings: list[str] = []
    total_pages = 0

    try:
        with pdfplumber.open(str(path)) as pdf, fitz.open(str(path)) as render_pdf:
            total_pages = len(pdf.pages)
            for i, page in enumerate(pdf.pages):
                # Quick check: does this page have meaningful text?
                quick_text = page.extract_text() or ""
                char_count = len(quick_text.replace(" ", "").replace("\n", ""))

                if char_count >= _MIN_DIGITAL_CHARS:
                    # Digital page — full coordinate-aware extraction
                    page_text = _extract_page_digital(page)
                else:
                    # Scanned / image-only page — render only this page for OCR.
                    scanned_pages.append(i + 1)
                    ocr = _ocr_pdf_page(render_pdf, i)
                    page_text = ocr.text
                    if ocr.text:
                        ocr_scores.append(ocr.confidence)
                        ocr_engines.add(ocr.engine)
                    else:
                        warnings.append(
                            ocr.warning or f"Page {i + 1} requires OCR but no text was recognized."
                        )
                    logger.info(
                        "[extractor] page %d/%d scanned chars=%d ocr=%s ocr_chars=%d",
                        i + 1, total_pages, char_count, ocr.engine, len(page_text),
                    )

                page_texts.append(page_text)
    except Exception as exc:
        logger.error("[extractor] PDF open failed (%s): %s", path.name, exc)
        return ExtractionResult(
            text="",
            warnings=[f"PDF extraction failed: {exc}"],
            confidence=0.0,
        )

    raw_text = "\n\n".join(t for t in page_texts if t.strip())
    clean = _clean_text(raw_text, page_texts)

    scanned_count = len(scanned_pages)
    ocr_required = scanned_count > 0

    method = "pdf_digital"
    if scanned_count == total_pages:
        method = "pdf_scanned"
    elif scanned_count > 0:
        method = "pdf_mixed"

    page_scores = [1.0] * (total_pages - scanned_count) + ocr_scores
    confidence = sum(page_scores) / total_pages if total_pages else 0.0
    failed_ocr_pages = scanned_count - len(ocr_scores)
    if failed_ocr_pages:
        warnings.append(
            f"OCR could not extract text from {failed_ocr_pages} scanned page(s): {scanned_pages}."
        )
    if ocr_engines:
        method = f"{method}_{'+'.join(sorted(ocr_engines))}"

    return ExtractionResult(
        text=clean,
        page_count=total_pages,
        scanned_page_count=scanned_count,
        ocr_required=ocr_required,
        extraction_method=method,
        warnings=warnings,
        confidence=confidence,
    )


# ---------------------------------------------------------------------------
# DOCX extraction
# ---------------------------------------------------------------------------

def _extract_docx_media(path: Path) -> tuple[list[str], list[float], list[str], set[str]]:
    """OCR substantial raster images embedded in a DOCX (skip icons/logos)."""
    from PIL import Image

    texts: list[str] = []
    scores: list[float] = []
    warnings: list[str] = []
    engines: set[str] = set()
    supported = {".png", ".jpg", ".jpeg", ".tif", ".tiff", ".bmp", ".webp"}

    try:
        with zipfile.ZipFile(path) as archive:
            media_names = [
                name for name in archive.namelist()
                if name.startswith("word/media/") and Path(name).suffix.lower() in supported
            ]
            for media_name in media_names:
                suffix = Path(media_name).suffix.lower()
                with tempfile.NamedTemporaryFile(suffix=suffix, delete=False) as tmp:
                    tmp.write(archive.read(media_name))
                    media_path = Path(tmp.name)
                try:
                    with Image.open(media_path) as image:
                        width, height = image.size
                    # CV screenshots/scans are substantial; tiny assets are usually icons.
                    if width < 400 or height < 200 or width * height < 150_000:
                        continue
                    result = _extract_image(media_path)
                    if result.text:
                        texts.append(result.text)
                        scores.append(result.confidence)
                        engines.add(result.extraction_method.removeprefix("image_"))
                    elif result.warnings:
                        warnings.extend(result.warnings)
                finally:
                    media_path.unlink(missing_ok=True)
    except Exception as exc:
        warnings.append(f"Embedded DOCX image extraction failed: {exc}")

    return texts, scores, warnings, engines


def _extract_docx(path: Path) -> ExtractionResult:
    """Extract text from a .docx file using python-docx."""
    from docx import Document as DocxDocument  # lazy import

    parts: list[str] = []
    try:
        doc = DocxDocument(str(path))

        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                parts.append(text)

        # Tables → "Header: Value" pairs (preserves meaning)
        for table in doc.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells if c.text.strip()]
                if cells:
                    if len(cells) == 2:
                        parts.append(f"{cells[0]}: {cells[1]}")
                    else:
                        parts.append(" | ".join(cells))

    except Exception as exc:
        logger.error("[extractor] DOCX extraction failed (%s): %s", path.name, exc)
        return ExtractionResult(
            text="",
            warnings=[f"DOCX extraction failed: {exc}"],
            confidence=0.0,
        )

    media_texts, media_scores, warnings, media_engines = _extract_docx_media(path)
    parts.extend(media_texts)
    raw = "\n".join(parts)
    clean = _clean_text(raw)
    has_regular_text = bool(parts[:-len(media_texts)] if media_texts else parts)
    confidence_scores = ([0.95] if has_regular_text else []) + media_scores
    confidence = (
        sum(confidence_scores) / len(confidence_scores)
        if confidence_scores
        else 0.0
    )
    method = "docx"
    if media_engines:
        method = f"docx+image_{'+'.join(sorted(media_engines))}"
    return ExtractionResult(
        text=clean,
        page_count=1,
        scanned_page_count=len(media_scores),
        ocr_required=bool(media_texts or warnings),
        extraction_method=method,
        warnings=warnings,
        confidence=confidence,
    )


# ---------------------------------------------------------------------------
# DOC extraction (old binary .doc via LibreOffice)
# ---------------------------------------------------------------------------

def _libreoffice_available() -> bool:
    """Check if LibreOffice is available on PATH."""
    for binary in ["libreoffice", "soffice"]:
        try:
            result = subprocess.run(
                [binary, "--version"],
                capture_output=True,
                timeout=10,
            )
            if result.returncode == 0:
                return True
        except (FileNotFoundError, subprocess.TimeoutExpired):
            continue
    return False


def _extract_doc(path: Path) -> ExtractionResult:
    """
    Extract text from old binary .doc by converting to DOCX via LibreOffice headless,
    then applying the DOCX extractor.
    """
    if not _libreoffice_available():
        logger.warning("[extractor] LibreOffice not found — cannot extract .doc: %s", path.name)
        return ExtractionResult(
            text="",
            warnings=[
                f"Cannot extract '{path.name}': LibreOffice is required to process old .doc files. "
                "Install LibreOffice and restart the server."
            ],
            confidence=0.0,
            extraction_method="doc_failed",
        )

    with tempfile.TemporaryDirectory() as tmp_dir:
        tmp_path = Path(tmp_dir)
        try:
            # Convert .doc → .docx in a temp directory
            result = subprocess.run(
                [
                    "libreoffice", "--headless", "--convert-to", "docx",
                    "--outdir", str(tmp_path),
                    str(path),
                ],
                capture_output=True,
                timeout=60,
            )
            if result.returncode != 0:
                err = result.stderr.decode(errors="replace")[:300]
                return ExtractionResult(
                    text="",
                    warnings=[f"LibreOffice conversion failed: {err}"],
                    confidence=0.0,
                    extraction_method="doc_failed",
                )

            # Find the converted file
            converted = list(tmp_path.glob("*.docx"))
            if not converted:
                return ExtractionResult(
                    text="",
                    warnings=["LibreOffice conversion produced no output file."],
                    confidence=0.0,
                    extraction_method="doc_failed",
                )

            result = _extract_docx(converted[0])
            result.extraction_method = "doc_via_libreoffice"
            return result

        except subprocess.TimeoutExpired:
            return ExtractionResult(
                text="",
                warnings=["LibreOffice conversion timed out (>60s)."],
                confidence=0.0,
                extraction_method="doc_failed",
            )
        except Exception as exc:
            return ExtractionResult(
                text="",
                warnings=[f"DOC extraction error: {exc}"],
                confidence=0.0,
                extraction_method="doc_failed",
            )


# ---------------------------------------------------------------------------
# Image extraction
# ---------------------------------------------------------------------------

def _preprocess_image_for_ocr(image_path: Path) -> Optional[Path]:
    """
    Apply OpenCV pre-processing pipeline to improve OCR accuracy:
    1. Grayscale
    2. Deskew (basic rotation detection)
    3. Adaptive threshold (binarize)
    4. Noise removal
    Returns path to processed image in a temp file, or None on failure.
    """
    try:
        import cv2
        import numpy as np

        img = cv2.imread(str(image_path))
        if img is None:
            return None

        # Grayscale
        gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)

        # Deskew using moments
        coords = np.column_stack(np.where(gray < 200))
        if len(coords) > 100:
            angle = cv2.minAreaRect(coords)[-1]
            if angle < -45:
                angle = 90 + angle
            if abs(angle) > 0.5:
                (h, w) = gray.shape
                center = (w // 2, h // 2)
                M = cv2.getRotationMatrix2D(center, angle, 1.0)
                gray = cv2.warpAffine(
                    gray, M, (w, h),
                    flags=cv2.INTER_CUBIC,
                    borderMode=cv2.BORDER_REPLICATE,
                )

        # Adaptive threshold — handles uneven lighting
        thresh = cv2.adaptiveThreshold(
            gray, 255,
            cv2.ADAPTIVE_THRESH_GAUSSIAN_C,
            cv2.THRESH_BINARY,
            11, 2,
        )

        # Noise removal (morphological opening)
        kernel = np.ones((1, 1), np.uint8)
        cleaned = cv2.morphologyEx(thresh, cv2.MORPH_OPEN, kernel)

        # Write to temp file
        tmp = tempfile.NamedTemporaryFile(suffix=".png", delete=False)
        tmp_path = Path(tmp.name)
        tmp.close()
        cv2.imwrite(str(tmp_path), cleaned)
        return tmp_path

    except Exception as exc:
        logger.warning("[extractor] image pre-processing failed: %s", exc)
        return None


def _extract_image(path: Path) -> ExtractionResult:
    """
    Extract text from image file (PNG, JPG, TIFF, BMP, WEBP).
    Pre-processes with OpenCV, then OCR.
    Uses PaddleOCR when configured and available, with Tesseract fallback in
    `auto` mode.
    """
    preprocessed = _preprocess_image_for_ocr(path)
    try:
        ocr = _run_ocr(path, preprocessed)
        return ExtractionResult(
            text=ocr.text,
            page_count=1,
            scanned_page_count=1,
            ocr_required=True,
            extraction_method=f"image_{ocr.engine}" if ocr.text else "image_no_ocr",
            warnings=[ocr.warning] if ocr.warning else [],
            confidence=ocr.confidence,
        )
    finally:
        if preprocessed:
            try:
                preprocessed.unlink(missing_ok=True)
            except OSError:
                pass


# ---------------------------------------------------------------------------
# Public API — single entry point
# ---------------------------------------------------------------------------

# MIME types → extractor routing
_PDF_MIMES = {"application/pdf"}
_DOCX_MIMES = {
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "application/zip",  # some DOCX files arrive with this MIME
}
_DOC_MIMES = {"application/msword"}
_IMAGE_MIMES = {
    "image/jpeg", "image/png", "image/tiff",
    "image/bmp", "image/webp", "image/gif",
}


def extract_document(path: Path, mime_type: str) -> ExtractionResult:
    """
    Main entry point. Route the file to the correct extractor based on MIME type.
    Falls back to extension if MIME is ambiguous.

    Args:
        path:       Absolute path to the file on disk.
        mime_type:  MIME type string (from documents.mime_type column).

    Returns:
        ExtractionResult with clean text, metadata, and confidence score.
    """
    if not path.exists():
        return ExtractionResult(
            text="",
            warnings=[f"File not found: {path}"],
            confidence=0.0,
        )

    mime = (mime_type or "").lower().strip()
    suffix = path.suffix.lower()

    logger.info(
        "[extractor] starting: file=%s mime=%s size=%d",
        path.name, mime, path.stat().st_size,
    )

    # MIME was established from validated file content at upload time. Only use
    # an extension fallback for legacy rows whose MIME is genuinely unknown.
    ambiguous_mime = mime in {"", "application/octet-stream"}
    if mime in _PDF_MIMES or (ambiguous_mime and suffix == ".pdf"):
        result = _extract_pdf(path)

    elif mime in _DOCX_MIMES or (ambiguous_mime and suffix == ".docx"):
        result = _extract_docx(path)

    elif mime in _DOC_MIMES or (ambiguous_mime and suffix == ".doc"):
        result = _extract_doc(path)

    elif mime in _IMAGE_MIMES or (
        ambiguous_mime
        and suffix in (".png", ".jpg", ".jpeg", ".tiff", ".tif", ".bmp", ".webp")
    ):
        result = _extract_image(path)

    else:
        logger.warning("[extractor] unknown MIME type '%s' for %s", mime, path.name)
        result = ExtractionResult(
            text="",
            warnings=[f"Unsupported file type: mime={mime} ext={suffix}"],
            confidence=0.0,
        )

    logger.info(
        "[extractor] done: file=%s method=%s pages=%d chars=%d scanned=%d confidence=%.2f",
        path.name,
        result.extraction_method,
        result.page_count,
        len(result.text),
        result.scanned_page_count,
        result.confidence,
    )

    return result
