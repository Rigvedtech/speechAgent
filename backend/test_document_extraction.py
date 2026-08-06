from __future__ import annotations

import io
import tempfile
import unittest
import uuid
import zipfile
from pathlib import Path
from unittest.mock import patch

import fitz
from docx import Document as DocxDocument
from fastapi import HTTPException
from PIL import Image, ImageDraw, ImageFont

from routers.uploads import _detect_mime, _require_ready_jd
from db.models import Candidate, Document, JobCandidateLink, JobPosting, User
from routers.parsing import _ensure_candidate_and_job_link
from services.structured_extractor import ParsedCV
from services.text_extractor import OcrResult, extract_document


def _image_bytes(text: str = "Candidate Resume") -> bytes:
    image = Image.new("RGB", (1200, 800), "white")
    try:
        font = ImageFont.truetype("arial.ttf", 64)
    except OSError:
        font = ImageFont.load_default()
    ImageDraw.Draw(image).text((80, 100), text, fill="black", font=font)
    output = io.BytesIO()
    image.save(output, format="PNG")
    return output.getvalue()


def _digital_pdf_bytes() -> bytes:
    pdf = fitz.open()
    page = pdf.new_page()
    page.insert_text((72, 72), "Pranay Sherkar Software Engineer Python FastAPI")
    data = pdf.tobytes()
    pdf.close()
    return data


def _scanned_pdf_bytes() -> bytes:
    pdf = fitz.open()
    page = pdf.new_page(width=612, height=792)
    page.insert_image(page.rect, stream=_image_bytes())
    data = pdf.tobytes()
    pdf.close()
    return data


def _docx_bytes(text: str = "Candidate Name\nPython Developer") -> bytes:
    document = DocxDocument()
    document.add_paragraph(text)
    output = io.BytesIO()
    document.save(output)
    return output.getvalue()


class UploadTypeDetectionTests(unittest.TestCase):
    def test_pdf_renamed_as_doc_uses_pdf_content(self) -> None:
        mime, extension = _detect_mime(_digital_pdf_bytes(), "resume.doc")
        self.assertEqual(mime, "application/pdf")
        self.assertEqual(extension, ".pdf")

    def test_image_renamed_as_pdf_uses_image_content(self) -> None:
        mime, extension = _detect_mime(_image_bytes(), "resume.pdf")
        self.assertEqual(mime, "image/png")
        self.assertEqual(extension, ".png")

    def test_docx_is_validated_as_word_package(self) -> None:
        mime, extension = _detect_mime(_docx_bytes(), "resume.bin")
        self.assertIn("wordprocessingml", mime)
        self.assertEqual(extension, ".docx")

    def test_plain_zip_is_not_accepted_as_docx(self) -> None:
        output = io.BytesIO()
        with zipfile.ZipFile(output, "w") as archive:
            archive.writestr("notes.txt", "not a Word file")
        with self.assertRaises(HTTPException):
            _detect_mime(output.getvalue(), "resume.docx")

    def test_extension_only_is_not_trusted(self) -> None:
        with self.assertRaises(HTTPException):
            _detect_mime(b"not really a PDF", "resume.pdf")


class ResumeUploadPrerequisiteTests(unittest.TestCase):
    def test_ready_processed_jd_allows_resume_upload(self) -> None:
        organization_id = uuid.uuid4()
        job_id = uuid.uuid4()
        document_id = uuid.uuid4()
        job = JobPosting(
            id=job_id,
            organization_id=organization_id,
            created_by=uuid.uuid4(),
            job_title="Backend Engineer",
            jd_document_id=document_id,
            pipeline_status="ready",
        )
        document = Document(
            id=document_id,
            organization_id=organization_id,
            job_posting_id=job_id,
            document_type="jd",
            upload_status="ready",
            extracted_text="Backend engineer job description",
        )

        class Session:
            def get(self, _model, _object_id):
                return document

        self.assertIs(_require_ready_jd(Session(), job), document)

    def test_job_without_processed_jd_rejects_resume_upload(self) -> None:
        job = JobPosting(
            id=uuid.uuid4(),
            organization_id=uuid.uuid4(),
            created_by=uuid.uuid4(),
            job_title="Backend Engineer",
            pipeline_status="pending",
        )

        class Session:
            def get(self, _model, _object_id):
                return None

        with self.assertRaises(HTTPException) as raised:
            _require_ready_jd(Session(), job)
        self.assertEqual(raised.exception.status_code, 409)

    def test_existing_extracted_jd_text_allows_legacy_pending_job(self) -> None:
        organization_id = uuid.uuid4()
        job_id = uuid.uuid4()
        document_id = uuid.uuid4()
        job = JobPosting(
            id=job_id,
            organization_id=organization_id,
            created_by=uuid.uuid4(),
            job_title="Imported Engineer",
            jd_document_id=document_id,
            jd_text="A complete imported job description " * 5,
            pipeline_status="pending",
        )
        document = Document(
            id=document_id,
            organization_id=organization_id,
            job_posting_id=job_id,
            document_type="jd",
            upload_status="ready",
            extracted_text=job.jd_text,
        )

        class Session:
            def get(self, _model, _object_id):
                return document

        self.assertIs(_require_ready_jd(Session(), job), document)

    def test_stored_jd_text_without_document_allows_resume_upload(self) -> None:
        job = JobPosting(
            id=uuid.uuid4(),
            organization_id=uuid.uuid4(),
            created_by=uuid.uuid4(),
            job_title="Saved Manual Job",
            jd_text="A complete manually entered job description " * 5,
            pipeline_status="pending",
        )

        class Session:
            def get(self, _model, _object_id):
                return None

        self.assertIsNone(_require_ready_jd(Session(), job))


class ExtractionTests(unittest.TestCase):
    def _temporary_file(self, data: bytes, suffix: str) -> tempfile.NamedTemporaryFile:
        file = tempfile.NamedTemporaryFile(suffix=suffix, delete=False)
        file.write(data)
        file.close()
        self.addCleanup(lambda: Path(file.name).unlink(missing_ok=True))
        return file

    def test_digital_pdf_extracts_without_ocr(self) -> None:
        file = self._temporary_file(_digital_pdf_bytes(), ".pdf")
        result = extract_document(Path(file.name), "application/pdf")
        self.assertIn("Pranay Sherkar", result.text)
        self.assertEqual(result.extraction_method, "pdf_digital")
        self.assertEqual(result.scanned_page_count, 0)

    def test_scanned_pdf_page_runs_ocr(self) -> None:
        file = self._temporary_file(_scanned_pdf_bytes(), ".pdf")
        mocked_ocr = OcrResult(
            text="Candidate Resume Python Developer",
            engine="test",
            confidence=0.91,
        )
        with patch("services.text_extractor._run_ocr", return_value=mocked_ocr):
            result = extract_document(Path(file.name), "application/pdf")
        self.assertIn("Python Developer", result.text)
        self.assertEqual(result.scanned_page_count, 1)
        self.assertEqual(result.extraction_method, "pdf_scanned_test")
        self.assertAlmostEqual(result.confidence, 0.91)

    def test_installed_tesseract_extracts_scanned_pdf(self) -> None:
        file = self._temporary_file(_scanned_pdf_bytes(), ".pdf")
        with patch("config.DOCUMENT_OCR_ENGINE", "tesseract"):
            result = extract_document(Path(file.name), "application/pdf")
        self.assertIn("CANDIDATE", result.text.upper())
        self.assertEqual(result.extraction_method, "pdf_scanned_tesseract")

    def test_docx_text_extraction(self) -> None:
        file = self._temporary_file(_docx_bytes(), ".docx")
        result = extract_document(
            Path(file.name),
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
        self.assertIn("Python Developer", result.text)
        self.assertEqual(result.extraction_method, "docx")

    def test_content_mime_wins_over_wrong_extension(self) -> None:
        file = self._temporary_file(_image_bytes(), ".doc")
        mocked_ocr = OcrResult(text="Image Resume", engine="test", confidence=0.9)
        with patch("services.text_extractor._run_ocr", return_value=mocked_ocr):
            result = extract_document(Path(file.name), "image/png")
        self.assertEqual(result.extraction_method, "image_test")
        self.assertIn("Image Resume", result.text)

    def test_installed_tesseract_extracts_image_text(self) -> None:
        file = self._temporary_file(
            _image_bytes("PRANAY SHERKAR\nPYTHON DEVELOPER"),
            ".png",
        )
        with patch("config.DOCUMENT_OCR_ENGINE", "tesseract"):
            result = extract_document(Path(file.name), "image/png")
        self.assertIn("PRANAY", result.text.upper())
        self.assertIn("PYTHON", result.text.upper())
        self.assertEqual(result.extraction_method, "image_tesseract")


class _ScalarResult:
    def __init__(self, value):
        self.value = value

    def first(self):
        return self.value

    def all(self):
        return self.value


class _FakeSession:
    def __init__(self, scalar_values):
        self.scalar_values = iter(scalar_values)
        self.added = []
        self.candidate = None

    def scalars(self, _statement):
        return _ScalarResult(next(self.scalar_values))

    def add(self, value):
        self.added.append(value)
        if isinstance(value, Candidate):
            self.candidate = value

    def flush(self):
        if self.candidate is not None and self.candidate.id is None:
            self.candidate.id = uuid.uuid4()

    def get(self, model, object_id):
        if model is Candidate and self.candidate and self.candidate.id == object_id:
            return self.candidate
        return None


class CandidateLinkingTests(unittest.TestCase):
    def test_parsed_bulk_cv_creates_candidate_and_job_link(self) -> None:
        organization_id = uuid.uuid4()
        user_id = uuid.uuid4()
        job_id = uuid.uuid4()
        doc_id = uuid.uuid4()
        actor = User(
            id=user_id,
            organization_id=organization_id,
            full_name="Recruiter",
            email="recruiter@example.com",
            role="recruiter",
        )
        document = Document(
            id=doc_id,
            organization_id=organization_id,
            uploaded_by=user_id,
            job_posting_id=job_id,
            document_type="cv",
            source="bulk_upload",
            original_filename="pranay_resume.pdf",
            extracted_text="Pranay Sherkar Python FastAPI",
            upload_status="ready",
        )
        parsed = ParsedCV(
            full_name="Pranay Sherkar",
            email="PRANAY@EXAMPLE.COM",
            phone="+91 9999999999",
            current_title="Python Developer",
            skills=["Python", "FastAPI"],
            domain_tags=["backend"],
        )
        # Candidate lookup, link lookup, then matching upload batch items.
        session = _FakeSession([None, None, []])

        candidate = _ensure_candidate_and_job_link(session, document, parsed, actor)

        self.assertEqual(candidate.email, "pranay@example.com")
        self.assertEqual(candidate.primary_cv_document_id, doc_id)
        self.assertEqual(document.candidate_id, candidate.id)
        links = [value for value in session.added if isinstance(value, JobCandidateLink)]
        self.assertEqual(len(links), 1)
        self.assertEqual(links[0].job_posting_id, job_id)
        self.assertEqual(links[0].cv_document_id, doc_id)


if __name__ == "__main__":
    unittest.main()
